import sqlite3
from PyQt5.QtWidgets import (
    QMainWindow, QPushButton, QVBoxLayout, QWidget,
    QLabel, QMessageBox, QInputDialog
)
from PyQt5.QtGui import QFont
import docx
from docx.shared import Pt
from Workout import Workout
from Exercise import Exercise


class Main(QMainWindow):
    """Главный класс приложения. Управление интерфейсом и логикой"""
    def __init__(self):
        """Инициализация главного окна и подключение к базе данных"""
        super().__init__()
        self.initUI()
        self.init_db()

    def initUI(self):
        """Создание элементов пользовательского интерфейса"""
        self.setWindowTitle("Дневник тренировок Сергазиева Руслана")
        self.setGeometry(100, 100, 400, 150)

        self.layout = QVBoxLayout()

        self.info_label = QLabel("Дневник тренировок:", self)
        self.info_label.setFont(QFont("Times New Roman", 14))
        self.layout.addWidget(self.info_label)

        self.save_workout_button = QPushButton("Сохранить тренировку", self)
        self.save_workout_button.clicked.connect(self.save_workout)
        self.layout.addWidget(self.save_workout_button)

        self.admin_button = QPushButton("Режим администратора", self)
        self.admin_button.clicked.connect(self.admin_mode)
        self.layout.addWidget(self.admin_button)

        self.exit_button = QPushButton("Выход", self)
        self.exit_button.clicked.connect(self.close)
        self.layout.addWidget(self.exit_button)

        container = QWidget()
        container.setLayout(self.layout)
        self.setCentralWidget(container)

    def init_db(self):
        """Инициализация таблиц базы данных, если они отсутствуют"""
        self.conn = sqlite3.connect("workouts.db")
        self.cursor = self.conn.cursor()
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS workouts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                date TEXT NOT NULL,
                details TEXT NOT NULL
            )
        """)
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                password TEXT NOT NULL
            )
        """)
        self.conn.commit()

    def save_workout(self):
        """Сохранение тренировки в базу данных и генерация отчёта"""
        login, login_confirmed = self.get_text_input("Введите логин", "авторизация")
        if not login_confirmed:
            return
        password, password_confirmed = self.get_text_input("Введите пароль", "авторизация")
        if not password_confirmed:
            return

        self.cursor.execute("SELECT password FROM users WHERE username = ?", (login,))
        result = self.cursor.fetchone()

        if result and result[0] == password:
            with open('new_workout.txt', 'r', encoding='utf-8') as file_in:
                date = file_in.readline().split()
                weight, height = file_in.readline().split()
                workout = Workout(f"{date[0]}-{date[1]}-{date[2]}", weight, height)

                for line in file_in:
                    line = line.strip()
                    parts = line.rsplit(maxsplit=3)
                    if len(parts) != 4:
                        continue
                    name, sets, reps, weight = parts[0], int(parts[1]), int(parts[2]), float(parts[3])
                    exercise = Exercise(name, sets, reps, weight)
                    workout.add_exercise(exercise)

                document = docx.Document()
                style = document.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                document.add_paragraph(workout.date_to_all()).style = style
                document.add_paragraph(workout.to_out_exercises()).style = style
                document.save(f"D:/{'-'.join(date) + ' ' + login}.docx")

                with open('last_workout.txt', 'w', encoding='utf-8') as file_out_last:
                    file_out_last.write(workout.date_to_last())
                    file_out_last.writelines(workout.to_out_exercises())

                with open('workouts_list.txt', 'a', encoding='utf-8') as file_out:
                    file_out.write(workout.date_to_all())
                    file_out.writelines(workout.to_out_exercises())

                self.cursor.execute(
                    "INSERT INTO workouts (username, date, details) VALUES (?, ?, ?)",
                    (login, workout.date, workout.to_out_exercises())
                )
                self.conn.commit()

                QMessageBox.information(self, "успешно", "Тренировка успешно сохранена")
        else:
            QMessageBox.warning(self, "ошибка", "Несуществующий пользователь или неверный пароль")

    def admin_mode(self):
        """Открытие режима администратора для управления пользователями и данными"""
        login, login_confirmed = self.get_text_input("Введите логин администратора", "админ-панель")
        if not login_confirmed or login != "admin":
            QMessageBox.warning(self, "ошибка", "неверный логин администратора")
            return

        password, password_confirmed = self.get_text_input("Введите пароль администратора", "админ-панель")
        if not password_confirmed or password != "admin-password":
            QMessageBox.warning(self, "ошибка", "неверный пароль администратора")
            return

        admin_choice, check = QInputDialog.getItem(
            self, "админ-панель", "Выберите действие",
            ["Добавить пользователя", "Удалить пользователя", "Очистить список тренировок"],
            0, False
        )

        if not check:
            return

        if admin_choice == "Добавить пользователя":
            self.add_user()
        elif admin_choice == "Удалить пользователя":
            self.remove_user()
        elif admin_choice == "Очистить список тренировок":
            self.clear_workouts()

    def add_user(self):
        """Добавление нового пользователя в базу данных"""
        login, login_confirmed = self.get_text_input("Введите логин", "добавить пользователя")
        if not login_confirmed:
            return
        password, password_confirmed = self.get_text_input("Введите пароль", "добавить пользователя")
        if not password_confirmed:
            return

        try:
            self.cursor.execute("INSERT INTO users (username, password) VALUES (?, ?)", (login, password))
            self.conn.commit()
            QMessageBox.information(self, "успешно", "Пользователь добавлен")
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "ошибка", "Пользователь с таким логином уже существует")

    def remove_user(self):
        """Удаление пользователя из базы данных"""
        login, login_confirmed = self.get_text_input("Введите логин пользователя для удаления", "удаление")
        if not login_confirmed:
            return

        self.cursor.execute("DELETE FROM users WHERE username = ?", (login,))
        self.conn.commit()

        QMessageBox.information(self, "успешно", "Пользователь удалён")

    def clear_workouts(self):
        """Очистка всех записей тренировок из текстового файла"""
        with open('workouts_list.txt', 'w', encoding='utf-8') as file:
            file.write("")
        QMessageBox.information(self, "успешно", "Список тренировок очищен")

    def get_text_input(self, message, title):
        """Открывает диалоговое окно для ввода текста."""
        return QInputDialog.getText(self, title, message)
